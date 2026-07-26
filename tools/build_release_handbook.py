from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "HariSmruti_Android_iOS_Release_and_Shorebird_Handbook.docx"

NAVY = RGBColor(25, 55, 85)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(95, 105, 115)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GOLD = "FFF4D6"
PALE_RED = "FDECEC"
WHITE = RGBColor(255, 255, 255)
CURRENT_NUM_ID = None


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=11, bold=False, color=None, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix):])
        set_run(r)
    else:
        set_run(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text))
    return p


def add_number(doc, text):
    global CURRENT_NUM_ID
    if CURRENT_NUM_ID is None:
        start_numbering(doc)
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(CURRENT_NUM_ID))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text))
    return p


def start_numbering(doc):
    global CURRENT_NUM_ID
    numbering = doc.part.numbering_part.element
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    CURRENT_NUM_ID = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(CURRENT_NUM_ID))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), "7")
    num.append(abstract_num_id)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    for index, line in enumerate(text.splitlines()):
        if index:
            p.add_run().add_break()
        set_run(p.add_run(line), size=9, font="Consolas", color=RGBColor(36, 44, 52))
    return p


def add_callout(doc, label, text, kind="info"):
    fill = {"info": LIGHT_BLUE, "warning": PALE_GOLD, "danger": PALE_RED}[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(f"{label}: "), bold=True, color=NAVY)
    set_run(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(header), size=9.5, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(value), size=9.2)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def new_page(doc, title, subtitle=None):
    doc.add_page_break()
    heading(doc, title, 1)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        set_run(p.add_run(subtitle), size=10.5, italic=True, color=MUTED)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.38)
section.footer_distance = Inches(0.42)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header.add_run("HARI SMRUTI  |  RELEASE OPERATIONS"), size=8.5, bold=True, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(footer.add_run("Internal handbook - keep credentials in a secure vault"), size=8.5, color=MUTED)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after = Pt(10)
set_run(p.add_run("RELEASE & CODE PUSH"), size=12, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
set_run(p.add_run("Android and iOS\nOperations Handbook"), size=30, bold=True, color=NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(24)
set_run(
    p.add_run("Play Store, TestFlight, Shorebird releases, patches, signing, testing, rollback, and recovery"),
    size=13,
    color=MUTED,
)
add_table(
    doc,
    ["Project", "Current production baseline"],
    [
        ("Application", "HariPrabodham Smruti"),
        ("Package / bundle ID", "org.hp.harismruti"),
        ("Shorebird app ID", "c46e42b3-37fe-4c81-a659-01d4f4c51bd6"),
        ("Next store baseline", "1.0.0+2"),
        ("Platforms", "Android and iOS"),
    ],
    [2500, 6860],
)
add_callout(
    doc,
    "Security",
    "This handbook intentionally contains no passwords, private keys, API tokens, or Apple credentials.",
    "warning",
)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(44)
set_run(p.add_run("Prepared 25 July 2026  |  Review CLI help and store policy before every major release"), size=9.5, color=MUTED)

# Quick decision page
new_page(doc, "1. Start Here: Choose Patch or Store Release", "Use this decision before running any Shorebird command.")
add_table(
    doc,
    ["Change type", "Shorebird patch?", "Required action"],
    [
        ("Dart UI, API logic, GetX state, validation", "Usually yes", "Dry-run, staging patch, test, then stable"),
        ("Pure Dart package already compatible", "Usually yes", "Dry-run and inspect warnings"),
        ("Kotlin, Java, Swift, Objective-C", "No", "Increase build number; new store release"),
        ("New native Flutter plugin", "No", "New store release before using plugin code"),
        ("Android/iOS permissions or entitlements", "No", "New store release"),
        ("New/changed image, font, bundled asset", "No", "New store release"),
        ("Flutter engine or Shorebird Flutter version", "No", "New store release"),
        ("Backend-only content/configuration", "Not needed", "Deploy backend/config normally"),
    ],
    [3500, 1800, 4060],
)
add_callout(doc, "Golden rule", "If Shorebird reports native or asset differences, stop and investigate. Do not routinely use allow-native-diffs or allow-asset-diffs.", "danger")
heading(doc, "Release identities", 2)
add_bullet(doc, "Version name is user-facing, for example 1.0.0.")
add_bullet(doc, "Build number/version code must increase for every Play Store or App Store Connect upload.")
add_bullet(doc, "A Shorebird patch keeps the same release identity, for example 1.0.0+2 Patch 1.")
add_bullet(doc, "Android and iOS may share a version number, but Shorebird releases and patches are platform-specific.")
heading(doc, "Pre-flight checklist", 2)
for item in (
    "Freeze the intended source state; do not patch while unrelated files are changing.",
    "Review git status and diffs. Preserve user work and commit the exact release source.",
    "Run flutter pub get, flutter analyze, and relevant tests.",
    "Confirm production API/Firebase configuration and store version gate.",
    "Confirm signing credentials are available without placing secrets in Git.",
    "Back up the final AAB/IPA and record its checksum.",
):
    add_bullet(doc, item)

# Android
new_page(doc, "2. Android: First Play Store Release", "Use Windows, macOS, or Linux. This project currently uses a Windows Shorebird installation.")
heading(doc, "Signing files", 2)
add_bullet(doc, r"Keystore: android\upload-keystore.jks (ignored by Git).")
add_bullet(doc, r"Properties: android\key.properties (ignored by Git).")
add_bullet(doc, "Store both files and their passwords in an encrypted vault. Never commit or email them.")
add_bullet(doc, "Upload alias for this project: upload.")
add_callout(doc, "Do not lose the key", "After Play publishing, future updates must remain compatible with the registered upload key. Configure Play App Signing and preserve recovery access.", "warning")
heading(doc, "Build the Shorebird-backed AAB", 2)
add_code(doc, r"""cd D:\Projects\new_hari\Harismruti_App
C:\Users\admin\.shorebird\bin\shorebird.ps1 doctor
C:\Users\admin\.shorebird\bin\shorebird.ps1 account whoami
C:\Users\admin\.shorebird\bin\shorebird.ps1 release android""")
start_numbering(doc)
add_number(doc, "Confirm the displayed app ID, version, platform, and Flutter version.")
add_number(doc, "Answer Yes only after the pre-flight checklist passes.")
add_number(doc, r"Copy build\app\outputs\bundle\release\app-release.aab to a versioned filename immediately.")
add_number(doc, "Record SHA-256 and retain the exact artifact; later patch builds overwrite app-release.aab.")
heading(doc, "Google Play Console", 2)
start_numbering(doc)
for item in (
    "Create the app listing with package org.hp.harismruti.",
    "Complete App content, privacy policy, data safety, content rating, target audience, and store listing.",
    "Enable Play App Signing and upload the Shorebird-generated AAB to Internal testing first.",
    "Resolve every blocking Play Console warning; test installation from the Play Store link.",
    "Promote Internal -> Closed/Open testing -> Production only after acceptance testing.",
):
    add_number(doc, item)

new_page(doc, "3. Android: Testing and Future Releases")
heading(doc, "Local preview before Play", 2)
add_code(doc, r"""C:\Users\admin\.shorebird\bin\shorebird.ps1 preview ^
  --platform android --release-version 1.0.0+2 --device-id DEVICE_ID ^
  --ks android\upload-keystore.jks --ks-pass file:PATH_TO_PASSWORD_FILE ^
  --ks-key-pass file:PATH_TO_KEY_PASSWORD_FILE --ks-key-alias upload""")
add_callout(doc, "Signature mismatch", "If a debug-signed copy is installed, Android will reject the production-signed preview. Uninstalling the old app clears only that app's local device data.", "warning")
heading(doc, "Every later Play Store release", 2)
start_numbering(doc)
for item in (
    "Increase pubspec.yaml build number: 1.0.0+3, then +4, and so on.",
    "Use a new version name when product semantics require it, for example 1.1.0+10.",
    "Run a fresh shorebird release android. Do not use flutter build appbundle for a patchable production base.",
    "Upload the exact Shorebird AAB to a Play test track and preserve a versioned copy.",
    "Never reuse an existing Play version code.",
):
    add_number(doc, item)
heading(doc, "Common Android cases", 2)
add_table(
    doc,
    ["Case", "Response"],
    [
        ("Play says version code already used", "Increase build number and create a new Shorebird release."),
        ("Wrong signing certificate", "Verify key.properties and upload key; do not create random replacement keys."),
        ("Existing debug app blocks install", "Back up needed local data, uninstall debug app, install signed build."),
        ("Gradle/AGP/Kotlin warning", "Plan toolchain upgrade as a new store release; never patch native build changes."),
        ("AAB overwritten after patch", "Use preserved versioned copy or retrieve exact release artifact; do not upload patch-build AAB."),
    ],
    [3100, 6260],
)

# iOS
new_page(doc, "4. iOS: TestFlight and App Store", "iOS archives and IPAs must be built on macOS with Xcode.")
heading(doc, "Current project configuration", 2)
add_table(
    doc,
    ["Setting", "Value"],
    [
        ("Bundle ID", "org.hp.harismruti"),
        ("Apple Team ID", "459224PK3J"),
        ("Minimum iOS", "15.5"),
        ("Signing", "Automatic in Xcode project"),
        ("Firebase", "GoogleService-Info.plist present"),
        ("Background mode", "Remote notifications enabled"),
        ("Next baseline", "1.0.0+2"),
    ],
    [2700, 6660],
)
heading(doc, "Mac prerequisites", 2)
for item in (
    "Current Xcode and accepted license; matching macOS support.",
    "Paid Apple Developer membership and access to Team 459224PK3J.",
    "App ID and App Store Connect record for org.hp.harismruti.",
    "Distribution certificate and App Store provisioning access, preferably automatic signing.",
    "CocoaPods and Shorebird CLI authenticated to the correct accounts.",
):
    add_bullet(doc, item)
heading(doc, "Create the TestFlight build", 2)
add_code(doc, """cd /path/to/Harismruti_App
flutter pub get
cd ios && pod install && cd ..
shorebird doctor
shorebird account whoami
shorebird release ios --export-method app-store""")
start_numbering(doc)
add_number(doc, "Confirm bundle ID, team, version, build number, and Shorebird Flutter version.")
add_number(doc, "Upload the generated IPA through Xcode Organizer or Apple Transporter.")
add_number(doc, "Wait for App Store Connect processing and answer export-compliance questions.")
add_number(doc, "Assign the build to Internal TestFlight testers first.")
add_number(doc, "External TestFlight testing requires Beta App Review.")

new_page(doc, "5. iOS: TestFlight Cases and Store Promotion")
add_table(
    doc,
    ["Case", "Action"],
    [
        ("No Mac available", "Use a trusted macOS CI provider; store Apple and Shorebird credentials as encrypted secrets."),
        ("Signing team missing", "Sign into Xcode and select Team 459224PK3J; verify account role."),
        ("Provisioning failure", "Open ios/Runner.xcworkspace in Xcode, enable automatic signing, resolve capabilities."),
        ("Build number already used", "Increase pubspec build number and create a new Shorebird iOS release."),
        ("TestFlight build processed", "Complete compliance, add testing notes, assign tester group."),
        ("External testers", "Submit build and testing information for Beta App Review."),
        ("Ready for App Store", "Create app version, complete metadata, select build, submit for App Review."),
        ("Native/asset change", "Create and distribute a new TestFlight/App Store release; do not patch."),
    ],
    [3000, 6360],
)
heading(doc, "iOS release cautions", 2)
add_bullet(doc, "Build the original release and all patches with compatible Xcode/Shorebird toolchains.")
add_bullet(doc, "Do not change entitlements, capabilities, plist permissions, native plugins, or assets in a patch.")
add_bullet(doc, "iOS patches may interpret changed Dart code; test performance-sensitive image and gallery paths.")
add_bullet(doc, "Keep functionality consistent with the reviewed app purpose and Apple policy.")
add_callout(doc, "TestFlight is not local preview", "TestFlight always uses App Store Connect. For faster device checks, use shorebird preview on a Mac with a registered iPhone.", "info")

# Patch workflow
new_page(doc, "6. Shorebird Patch Workflow: Android and iOS")
heading(doc, "Safe staging-first procedure", 2)
start_numbering(doc)
for item in (
    "Start from the exact source and lockfile associated with the target store release.",
    "Make only the intended Dart change.",
    "Run analysis and tests.",
    "Run the platform patch with --dry-run.",
    "Stop on asset/native warnings and create a store release instead.",
    "Publish to staging; test on the same platform and release version.",
    "Verify first launch/download, full close, second launch/activation.",
    "Promote the tested patch to stable.",
    "Monitor crashes, API errors, support reports, and Shorebird console status.",
):
    add_number(doc, item)
heading(doc, "Commands", 2)
add_code(doc, """# Android
shorebird patch android --release-version 1.0.0+2 --track staging --dry-run
shorebird patch android --release-version 1.0.0+2 --track staging

# iOS (run on Mac)
shorebird patch ios --release-version 1.0.0+2 --track staging --dry-run
shorebird patch ios --release-version 1.0.0+2 --track staging

# Promote after validation
shorebird patches set-track --release 1.0.0+2 --patch PATCH_NUMBER --track stable""")
heading(doc, "Expected device behavior", 2)
add_bullet(doc, "The updater checks over Wi-Fi/mobile data; production patches do not require USB.")
add_bullet(doc, "A patch downloads while the app runs and activates on a full process restart.")
add_bullet(doc, "The app's runtime service also checks after first frame and on foreground resume, throttled to 15 minutes.")
add_bullet(doc, "The profile/store version remains 1.0.0; Shorebird tracks patch numbers separately.")

new_page(doc, "7. Patch Recovery, Rollout, and Emergency Cases")
add_table(
    doc,
    ["Situation", "Immediate response", "Long-term response"],
    [
        ("Patch fails dry-run", "Do not upload", "Remove incompatible changes or create store release"),
        ("Patch crashes staging", "Keep it off stable", "Fix Dart code and create a new staging patch"),
        ("Bad patch reached stable", "Move/disable affected track in Shorebird console; publish corrected patch", "Document incident and add regression test"),
        ("User has old store version", "Patch that exact supported release if safe", "Use /app-version/check to require store upgrade"),
        ("Patch not downloading", "Check internet, release version, platform, track, and full restart", "Inspect Shorebird console/device logs"),
        ("New asset required urgently", "Do not force asset diff", "Ship new Play/TestFlight release"),
        ("Backend contract changed", "Keep API backward-compatible", "Coordinate minimum app version and deprecation window"),
    ],
    [2300, 3300, 3760],
)
add_callout(doc, "Rollback reality", "A corrected patch is usually safer than trying to hot-restart or manipulate files on devices. Preserve a known-good source tag for every release.", "warning")
heading(doc, "Store version gate and Shorebird", 2)
add_body(doc, "Keep both systems. Shorebird updates Dart code within an installed native release. The existing /app-version/check endpoint requires a Play/App Store upgrade when native code, assets, permissions, plugins, or the engine change.")

# CI and records
new_page(doc, "8. Automation, Secrets, and Release Records")
heading(doc, "Secrets that belong in a vault/CI secret store", 2)
for item in (
    "SHOREBIRD_TOKEN generated from Shorebird account API keys.",
    "Android upload keystore, alias, store password, and key password.",
    "App Store Connect API issuer ID, key ID, and private .p8 key.",
    "Apple certificate/provisioning material when automatic CI signing is not used.",
    "Production API keys and dart-define files.",
):
    add_bullet(doc, item)
add_callout(doc, "Never commit", "key.properties, .jks/.keystore, .p8, certificates, provisioning profiles, password files, or exported CI environment files.", "danger")
heading(doc, "Record for every store release", 2)
add_table(
    doc,
    ["Field", "Example"],
    [
        ("Git commit/tag", "release/1.0.0+2"),
        ("Platform", "Android or iOS"),
        ("Store version", "1.0.0+2"),
        ("Shorebird Flutter", "3.44.7 / revision recorded by CLI"),
        ("Build machine", "Windows Android builder or named Mac runner"),
        ("Artifact checksum", "SHA-256 of preserved AAB/IPA"),
        ("Signing identity", "Certificate fingerprint; never password"),
        ("Store track", "Play Internal / TestFlight Internal / Production"),
        ("Known patches", "Patch numbers and tracks"),
    ],
    [3000, 6360],
)
heading(doc, "Suggested CI gates", 2)
add_bullet(doc, "Formatting, static analysis, unit/widget tests, and release-mode build.")
add_bullet(doc, "Patch dry-run on intended release version before upload.")
add_bullet(doc, "Manual approval for store release, stable patch promotion, and production rollout.")
add_bullet(doc, "Artifact checksum, changelog, and release record attached to the build.")

# Troubleshooting
new_page(doc, "9. Troubleshooting Quick Reference")
add_table(
    doc,
    ["Message / symptom", "Meaning and fix"],
    [
        ("shorebird not recognized", "CLI is not on PATH. Use full script path on Windows or add its bin directory to user PATH."),
        ("Failed to refresh credentials", "Run shorebird logout, then login and complete browser authorization."),
        ("Git long paths warning", "Enable Windows Git long paths with administrator policy when possible."),
        ("INSTALL_FAILED_UPDATE_INCOMPATIBLE", "Installed package uses another signature. Uninstall after acknowledging local data loss."),
        ("No connected devices", "Unlock device, enable USB debugging, approve computer, verify adb/flutter devices."),
        ("Updater unavailable", "Normal flutter run/build lacks Shorebird engine. Use shorebird release or preview."),
        ("Patch does not appear", "Confirm target release/track; launch to download, fully close, then launch again."),
        ("Native/asset diff warning", "Do not bypass by default. Create a new store release."),
        ("Kotlin incremental cache error", "Stop Gradle, clean build caches safely, then rebuild; verify final command status."),
        ("Gradle/AGP/Kotlin deprecation warning", "Schedule coordinated native toolchain upgrade in a new store release."),
        ("Play/TestFlight version already used", "Increment build number and create another Shorebird release."),
    ],
    [3300, 6060],
)
heading(doc, "Read-only diagnosis commands", 2)
add_code(doc, """flutter doctor -v
flutter devices
shorebird doctor
shorebird account whoami
shorebird releases list
shorebird patches list --release-version RELEASE_VERSION
git status --short
git diff --check""")

# Final checklists
new_page(doc, "10. Copy/Paste Checklists")
heading(doc, "Store release acceptance", 2)
for item in (
    "Version/build number increased and matches console upload.",
    "Exact source committed/tagged; lockfile retained.",
    "No analysis errors; required tests pass.",
    "Production API, Firebase, notifications, and minimum-version gate verified.",
    "Signing identity verified; secret files remain ignored.",
    "Built with shorebird release for the correct platform.",
    "AAB/IPA preserved under a versioned filename with checksum.",
    "Installed from Play test track or TestFlight and smoke-tested.",
    "Login, home/gallery, photo operations, camera, notifications, and update checks tested.",
    "Rollout and monitoring owner identified.",
):
    add_bullet(doc, f"☐ {item}")
heading(doc, "Patch acceptance", 2)
for item in (
    "Target platform and store release version confirmed.",
    "Only patchable Dart changes included.",
    "No new/changed native code, plugin requirement, permissions, entitlements, or assets.",
    "Dry-run reports no issues.",
    "Patch published to staging first.",
    "First launch downloads; full close; second launch activates.",
    "Critical user journeys tested on physical device.",
    "Patch promoted to stable only after approval.",
    "Monitoring confirms healthy rollout.",
):
    add_bullet(doc, f"☐ {item}")
heading(doc, "Authoritative references", 2)
for text in (
    "Shorebird documentation: https://docs.shorebird.dev/code-push/",
    "Google Play Console help: https://support.google.com/googleplay/android-developer/",
    "Apple TestFlight: https://developer.apple.com/testflight/",
    "App Store Connect help: https://developer.apple.com/help/app-store-connect/",
):
    add_bullet(doc, text)
add_callout(doc, "Operational principle", "Test privately, preserve exact artifacts, promote deliberately, and use a new store release whenever the native shell or bundled assets change.", "info")

doc.core_properties.title = "HariSmruti Android and iOS Release and Shorebird Handbook"
doc.core_properties.subject = "Play Store, TestFlight, and Shorebird operations"
doc.core_properties.author = "HariSmruti Release Engineering"
doc.core_properties.keywords = "Flutter, Shorebird, Android, iOS, Play Store, TestFlight"
doc.save(OUTPUT)
print(OUTPUT)
